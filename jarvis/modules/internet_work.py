from __future__ import annotations

import smtplib
from datetime import datetime
from email.message import EmailMessage
from pathlib import Path

import pandas as pd
from docx import Document

from jarvis.modules import base_actions as act
from jarvis.utils.logging_utils import action_guard

MODULE = "Internet & Work"


def aliases(ru: list[str], en: list[str]) -> list[str]:
    return ru + en


@action_guard(MODULE)
def google_search(*, core, transcript: str) -> None:
    q = transcript.replace("поиск", "").replace("гугл", "").replace("google", "").strip() or "python"
    act.open_url(f"https://www.google.com/search?q={q.replace(' ', '+')}")
    core.speak(f"Ищу в Google: {q}, сэр.")


@action_guard(MODULE)
def wiki_search(*, core, transcript: str) -> None:
    q = transcript.replace("википедия", "").replace("wiki", "").strip() or "искусственный интеллект"
    act.open_url(f"https://ru.wikipedia.org/w/index.php?search={q.replace(' ', '+')}")
    core.speak("Открываю результаты Википедии, сэр.")


@action_guard(MODULE)
def youtube_search(*, core, transcript: str) -> None:
    q = transcript.replace("ютуб", "").replace("youtube", "").replace("поиск", "").strip() or "lofi"
    act.open_url(f"https://www.youtube.com/results?search_query={q.replace(' ', '+')}")
    core.speak(f"Ищу видео на YouTube: {q}, сэр.")


@action_guard(MODULE)
def create_docx(*, core, transcript: str) -> None:
    target = Path.home() / "Documents" / f"jarvis_{datetime.now():%Y%m%d_%H%M%S}.docx"
    doc = Document()
    doc.add_heading("Документ JARVIS", level=1)
    doc.add_paragraph("Создано по голосовой команде.")
    doc.save(target)
    core.speak(f"Документ создан: {target.name}, сэр.")


@action_guard(MODULE)
def create_xlsx(*, core, transcript: str) -> None:
    target = Path.home() / "Documents" / f"jarvis_{datetime.now():%Y%m%d_%H%M%S}.xlsx"
    pd.DataFrame([{"задача": "создано jarvis", "время": datetime.now().isoformat()}]).to_excel(target, index=False)
    core.speak(f"Таблица создана: {target.name}, сэр.")


@action_guard(MODULE)
def send_email(*, core, transcript: str) -> None:
    sender = core.config.get("smtp_sender")
    password = core.config.get("smtp_password")
    recipient = core.config.get("smtp_recipient")
    server = core.config.get("smtp_server", "smtp.gmail.com")
    if not all([sender, password, recipient]):
        core.speak("Сэр, SMTP-настройки не заполнены.")
        return
    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = recipient
    msg["Subject"] = "Сообщение от JARVIS"
    msg.set_content("Письмо отправлено голосовой командой JARVIS.")
    with smtplib.SMTP_SSL(server, 465) as smtp:
        smtp.login(sender, password)
        smtp.send_message(msg)
    core.speak("Письмо отправлено, сэр.")


def register(registry) -> None:
    registry.register("google", aliases(["поиск в гугл", "найди в гугле", "гугл это", "поиск гугл", "найти в интернете", "найди в сети", "гугли"], ["google search", "search google", "find on google"]), MODULE, google_search)
    registry.register("wiki", aliases(["поиск в википедии", "википедия", "найди в вики", "открой статью", "вики поиск", "википедия поиск", "найти википедия"], ["search wikipedia", "wiki search", "open wikipedia"]), MODULE, wiki_search)
    registry.register("youtube", aliases(["поиск на ютубе", "найди видео", "ютуб поиск", "открой видео", "видео в ютуб", "найти ролик", "ищи на ютуб"], ["search youtube", "find video", "youtube lookup"]), MODULE, youtube_search)
    registry.register("close_tab", aliases(["закрой вкладку", "убери вкладку", "закрыть текущую", "вкладку закрыть", "закрой таб", "удали вкладку", "выход из вкладки"], ["close tab", "close browser tab", "exit tab"]), MODULE, lambda **k: act.hotkey("ctrl", "w"))
    registry.register("history", aliases(["открой историю", "история браузера", "покажи историю", "журнал посещений", "история вкладок", "история серфинга", "история"], ["open history", "browser history", "view history"]), MODULE, lambda **k: act.hotkey("ctrl", "h"))
    registry.register("incognito", aliases(["режим инкогнито", "открой инкогнито", "приватный режим", "секретный режим", "новое инкогнито", "приватное окно", "инкогнито окно"], ["open incognito", "private browsing", "secret mode"]), MODULE, lambda **k: act.hotkey("ctrl", "shift", "n"))
    registry.register("new_tab", aliases(["новая вкладка", "открой вкладку", "создай вкладку", "новый таб", "добавь вкладку", "сделай новую вкладку", "вкладка новая"], ["new tab", "open new tab", "create tab"]), MODULE, lambda **k: act.hotkey("ctrl", "t"))
    registry.register("refresh", aliases(["обнови страницу", "перезагрузи страницу", "рефреш", "обновление страницы", "перезагрузка вкладки", "обнови вкладку", "повтори загрузку"], ["refresh page", "reload page", "page reload"]), MODULE, lambda **k: act.press("f5"))
    registry.register("address_bar", aliases(["адресная строка", "фокус на адрес", "перейди в адрес", "выдели url", "строка адреса", "поле адреса", "адрес сайта"], ["focus address bar", "url bar", "go to address bar"]), MODULE, lambda **k: act.hotkey("ctrl", "l"))
    registry.register("create_doc", aliases(["создай документ", "новый docx", "создай ворд", "создать документ", "новый word файл", "документ ворд", "сделай документ"], ["create document", "new docx", "make word file"]), MODULE, create_docx)
    registry.register("create_sheet", aliases(["создай таблицу", "новый excel", "создать xlsx", "сделай spreadsheet", "эксель файл", "новая ведомость", "таблица эксель"], ["create spreadsheet", "new excel", "make xlsx"]), MODULE, create_xlsx)
    registry.register("send_email", aliases(["отправь письмо", "отправить email", "пошли почту", "отправь имейл", "направь письмо", "отправка почты", "письмо отправить"], ["send email", "email now", "mail message"]), MODULE, send_email)
    registry.register("calendar", aliases(["открой календарь", "расписание", "создай событие", "календарь google", "планировщик", "встреча в календаре", "управление календарем"], ["open calendar", "create calendar event", "manage calendar"]), MODULE, lambda **k: act.open_url("https://calendar.google.com"))
    registry.register("telegram", aliases(["открой телеграм", "сообщение телеграм", "телеграмм", "запусти telegram", "написать в телеграм", "telegram message", "телеграм"], ["send telegram message", "open telegram", "telegram automation"]), MODULE, lambda **k: act.launch("start telegram"))
    registry.register("discord", aliases(["открой дискорд", "сообщение дискорд", "дискорд", "запусти discord", "написать в дискорд", "discord message", "discord"], ["send discord message", "open discord", "discord automation"]), MODULE, lambda **k: act.launch("start discord"))
